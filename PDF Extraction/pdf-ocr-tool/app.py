"""
app.py - PDF OCR Tool Backend (v3)
Supports: sections, text cleanup, ToC scan, keyword scan, docx + zip export.
"""

import os, re, uuid, json, shutil, base64, zipfile, io
from pathlib import Path
from flask import Flask, request, jsonify, Response, send_from_directory, send_file

app = Flask(__name__, static_folder="static", static_url_path="")

BASE_DIR    = Path(__file__).parent
WORKSPACE   = BASE_DIR / "workspace"
WORKSPACE.mkdir(exist_ok=True)

# ── Dependencies ───────────────────────────────────────────────────────────────
try:
    import fitz
except ImportError:
    raise RuntimeError("Run: pip install pymupdf")

try:
    import pytesseract
    from PIL import Image, ImageEnhance, ImageFilter
except ImportError:
    raise RuntimeError("Run: pip install pytesseract pillow")

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# ── Text helpers ───────────────────────────────────────────────────────────────
DEVANAGARI_RE = re.compile(r'[\u0900-\u097F]')

def is_devanagari(text: str) -> bool:
    return bool(DEVANAGARI_RE.search(text))

def clean_ocr_text(raw: str) -> str:
    lines  = raw.split('\n')
    result = []
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            if result and result[-1] != '':
                result.append('')
            i += 1
            continue
        stripped = line.strip()
        if is_devanagari(stripped):
            if result and result[-1] != '':
                result.append('')
            result.append(stripped)
            result.append('')
            i += 1
            continue
        # English paragraph accumulation
        paragraph = stripped
        if paragraph.endswith('-'):
            paragraph = paragraph[:-1]
            i += 1
            if i < len(lines):
                nxt = lines[i].strip()
                if nxt and not is_devanagari(nxt):
                    paragraph += nxt
                    i += 1
                else:
                    result.append(paragraph)
                    continue
        else:
            i += 1
        while i < len(lines):
            nxt = lines[i].rstrip().strip()
            if not nxt: break
            if is_devanagari(nxt): break
            if paragraph.endswith('-'):
                paragraph = paragraph[:-1] + nxt
            else:
                paragraph = paragraph + ' ' + nxt
            i += 1
        result.append(paragraph)

    cleaned, prev_blank = [], False
    for ln in result:
        if ln == '':
            if not prev_blank:
                cleaned.append('')
            prev_blank = True
        else:
            cleaned.append(ln)
            prev_blank = False
    return '\n'.join(cleaned).strip()

def preprocess_image(img: Image.Image) -> Image.Image:
    img = img.convert("L")
    img = img.filter(ImageFilter.SHARPEN)
    return ImageEnhance.Contrast(img).enhance(2.0)

def _safe_dirname(name: str) -> str:
    return re.sub(r'[^\w\s-]', '', name).strip().replace(' ', '_')[:50] or "section"

def _ocr_page_image(pix, lang: str, psm: int = 6) -> str:
    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
    img = preprocess_image(img)
    return pytesseract.image_to_string(img, lang=lang, config=f"--oem 1 --psm {psm}")


# ── ToC parser ─────────────────────────────────────────────────────────────────

# Devanagari digit map → Arabic
_DEVA_DIGITS = str.maketrans('०१२३४५६७८९', '0123456789')

def _to_arabic(s: str) -> str:
    return s.translate(_DEVA_DIGITS)

# Strip common OCR noise chars from a section name
_NOISE_RE = re.compile(r'[|\\\[\]{}#@$%^&*_~`\'"<>]')
def _clean_name(s: str) -> str:
    s = _NOISE_RE.sub('', s)
    s = re.sub(r'\s{2,}', ' ', s)
    return s.strip(' .-–—')

def parse_toc(ocr_text: str, pdf_total: int, offset: int = 0) -> list:
    sections = []
    for line in ocr_text.split('\n'):
        line = line.strip()
        # Normalise Devanagari digits in the line for number extraction
        norm = _to_arabic(line)
        if not norm or len(norm) < 3:
            continue

        # Strategy 1: "Name ....... 123" (dots/dashes/spaces as leader)
        m = re.search(r'^(.+?)\s*[.·\-–—_|\s]{2,}\s*(\d{1,4})\s*$', norm)
        if m:
            name     = _clean_name(m.group(1))
            page_num = int(m.group(2)) + offset
            if len(name) > 1 and 0 < page_num <= pdf_total + abs(offset) + 50:
                sections.append({"name": name, "start": page_num})
            continue

        # Strategy 2: last whitespace-separated token is a number
        parts = norm.rsplit(None, 1)
        if len(parts) == 2 and parts[1].isdigit():
            page_num = int(parts[1]) + offset
            name     = _clean_name(parts[0])
            if len(name) > 1 and 0 < page_num <= pdf_total + abs(offset) + 50:
                sections.append({"name": name, "start": page_num})

    # Deduplicate & sort by start page
    seen, unique = set(), []
    for s in sorted(sections, key=lambda x: x['start']):
        if s['start'] not in seen:
            seen.add(s['start'])
            unique.append(s)

    # Assign end pages
    for i, sec in enumerate(unique):
        sec['end'] = unique[i+1]['start'] - 1 if i+1 < len(unique) else pdf_total
    return unique


# ── Docx builder ───────────────────────────────────────────────────────────────
def build_docx(sections_data: list) -> bytes:
    doc   = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    for section in sections_data:
        h = doc.add_heading(section['name'], level=1)
        h.runs[0].font.color.rgb = RGBColor(0x5B, 0x4C, 0xE0)
        for pi in section['pages']:
            lp = doc.add_paragraph()
            lr = lp.add_run(f"── Page {pi['page']} ──")
            lr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            lr.font.size = Pt(9); lr.font.italic = True
            for para_text in re.split(r'\n{2,}', pi.get('text','').strip()):
                para_text = para_text.strip()
                if not para_text: continue
                p = doc.add_paragraph()
                for li, ln in enumerate(para_text.split('\n')):
                    run = p.add_run(ln)
                    if is_devanagari(ln):
                        run.font.name = 'Noto Sans Devanagari'
                        run.font.size = Pt(13); run.bold = True
                    else:
                        run.font.name = 'Calibri'; run.font.size = Pt(11)
                    if li < len(para_text.split('\n')) - 1:
                        run.add_break()
        doc.add_paragraph(); doc.add_paragraph('─'*60); doc.add_paragraph()
    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.read()


# ── Routes ─────────────────────────────────────────────────────────────────────
@app.route("/")
def index():
    return send_from_directory("static", "index.html")

@app.route("/upload", methods=["POST"])
def upload():
    if "pdf" not in request.files:
        return jsonify({"error": "No file provided"}), 400
    f = request.files["pdf"]
    if not f.filename.lower().endswith(".pdf"):
        return jsonify({"error": "Only PDF files are supported"}), 400
    job_id   = str(uuid.uuid4())[:8]
    job_dir  = WORKSPACE / job_id
    job_dir.mkdir(parents=True, exist_ok=True)
    pdf_path = job_dir / "source.pdf"
    f.save(str(pdf_path))
    doc = fitz.open(str(pdf_path))
    count = len(doc); doc.close()
    return jsonify({"job_id": job_id, "filename": f.filename, "page_count": count})

@app.route("/thumbnail/<job_id>/<int:page_num>")
def thumbnail(job_id, page_num):
    img_path = WORKSPACE / job_id / "images" / f"page_{page_num}.png"
    if not img_path.exists():
        return jsonify({"error": "Not found"}), 404
    with open(img_path, "rb") as f:
        return jsonify({"image": base64.b64encode(f.read()).decode()})


# ── POST /scan-toc ─────────────────────────────────────────────────────────────
@app.route("/scan-toc", methods=["POST"])
def scan_toc():
    body      = request.get_json()
    job_id    = body.get("job_id")
    toc_start = int(body.get("toc_start", 1))
    toc_end   = int(body.get("toc_end", 5))
    lang      = body.get("lang", "eng+hin+san")
    offset    = int(body.get("offset", 0))

    pdf_path  = WORKSPACE / job_id / "source.pdf"
    if not pdf_path.exists():
        return jsonify({"error": "Job not found"}), 404

    doc   = fitz.open(str(pdf_path))
    total = len(doc)
    zoom  = 200 / 72
    mat   = fitz.Matrix(zoom, zoom)
    full_text = ""

    for page_num in range(toc_start, min(toc_end, total) + 1):
        page = doc.load_page(page_num - 1)
        pix   = page.get_pixmap(matrix=mat)
        text  = _ocr_page_image(pix, lang)
        full_text += text + "\n"

    doc.close()
    sections = parse_toc(full_text, total, offset)
    return jsonify({"sections": sections, "raw_text": full_text})


# ── POST /parse-toc-raw ────────────────────────────────────────────────────────
@app.route("/parse-toc-raw", methods=["POST"])
def parse_toc_raw():
    body     = request.get_json()
    job_id   = body.get("job_id")
    raw_text = body.get("raw_text", "")
    offset   = int(body.get("offset", 0))

    pdf_path = WORKSPACE / job_id / "source.pdf"
    if not pdf_path.exists():
        return jsonify({"error": "Job not found"}), 404

    doc = fitz.open(str(pdf_path))
    total = len(doc)
    doc.close()

    sections = parse_toc(raw_text, total, offset)
    return jsonify({"sections": sections})


# ── POST /scan-keyword  (SSE) ──────────────────────────────────────────────────
@app.route("/scan-keyword", methods=["POST"])
def scan_keyword():
    body       = request.get_json()
    job_id     = body.get("job_id")
    keyword    = body.get("keyword", "Chapter").strip()
    use_regex  = body.get("use_regex", False)
    scan_start = int(body.get("scan_start", 1))
    scan_end   = int(body.get("scan_end", 0))
    # Use the user's chosen language, NOT hardcoded 'eng'
    lang       = body.get("lang", "eng+hin+san")
    offset     = int(body.get("offset", 0))   # reserved for future use

    pdf_path   = WORKSPACE / job_id / "source.pdf"
    if not pdf_path.exists():
        return jsonify({"error": "Job not found"}), 404

    def generate():
        doc   = fitz.open(str(pdf_path))
        total = len(doc)
        end   = min(scan_end, total) if scan_end > 0 else total
        # Compile regex if enabled
        kw_pattern = None
        if use_regex:
            try:
                kw_pattern = re.compile(keyword, re.IGNORECASE)
            except Exception as e:
                yield f"data: {json.dumps({'type':'error','text': f'Invalid Regex: {e}'})}\n\n"
                return
        else:
            kw = keyword.lower()

        # 200 DPI, top 40% crop — better quality for scanned PDFs
        zoom = 200 / 72
        mat  = fitz.Matrix(zoom, zoom)

        boundaries = []

        for page_num in range(scan_start, end + 1):
            page = doc.load_page(page_num - 1)

            # OCR top 40% crop at 200 DPI (scanned PDFs)
            rect = page.rect
            clip = fitz.Rect(0, 0, rect.width, rect.height * 0.4)
            pix  = page.get_pixmap(matrix=mat, clip=clip)
            text = _ocr_page_image(pix, lang, psm=6)

            # Check keyword (case-insensitive, also check Devanagari-normalised)
            text_norm = _to_arabic(text)
            match_line = ""
            
            for ln in text_norm.split('\n'):
                if len(ln.strip()) > 2:
                    if use_regex:
                        if kw_pattern.search(ln):
                            match_line = _clean_name(ln)
                            break
                    else:
                        if kw in ln.lower():
                            match_line = _clean_name(ln)
                            break
            
            found = bool(match_line)
            if found:
                boundaries.append({"page": page_num, "text": match_line})

            done = page_num - scan_start + 1
            pct  = round(done / max(end - scan_start + 1, 1) * 100)
            yield f"data: {json.dumps({'type':'progress','page':page_num,'found':found,'match':match_line,'done':done,'total':end - scan_start + 1,'pct':pct})}\n\n"

        doc.close()

        # Build sections from boundaries
        sections = []
        for i, b in enumerate(boundaries):
            end_page = boundaries[i+1]['page'] - 1 if i+1 < len(boundaries) else total
            sections.append({
                "name":  b['text'][:80].strip(),
                "start": b['page'],
                "end":   end_page
            })

        yield f"data: {json.dumps({'type':'done','sections':sections,'total_found':len(boundaries)})}\n\n"

    return Response(generate(), mimetype="text/event-stream",
                    headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})


# ── POST /extract  (SSE) ───────────────────────────────────────────────────────
@app.route("/extract", methods=["POST"])
def extract():
    body     = request.get_json()
    job_id   = body.get("job_id")
    sections = body.get("sections", [])
    lang     = body.get("lang", "eng+hin+san")
    dpi      = int(body.get("dpi", 200))

    job_dir    = WORKSPACE / job_id
    pdf_path   = job_dir / "source.pdf"
    images_dir = job_dir / "images"
    images_dir.mkdir(exist_ok=True)

    if not pdf_path.exists():
        return jsonify({"error": "Job not found"}), 404

    total_pages = sum(max(0, s['end'] - s['start'] + 1) for s in sections)

    def generate():
        doc  = fitz.open(str(pdf_path))
        zoom = dpi / 72
        mat  = fitz.Matrix(zoom, zoom)
        done = 0
        all_data = []

        for si, section in enumerate(sections):
            name    = section.get('name', f'Section {si+1}')
            start   = int(section['start'])
            end     = int(section['end'])
            sec_dir = job_dir / _safe_dirname(name)
            sec_dir.mkdir(exist_ok=True)

            yield f"data: {json.dumps({'type':'section_start','section':name,'index':si+1,'total_sections':len(sections)})}\n\n"
            sec_pages = []

            for page_num in range(start, end + 1):
                if page_num < 1 or page_num > len(doc):
                    yield f"data: {json.dumps({'type':'error','section':name,'page':page_num,'text':f'Page {page_num} out of range.'})}\n\n"
                    done += 1
                    continue

                page     = doc.load_page(page_num - 1)
                pix      = page.get_pixmap(matrix=mat)
                img_path = images_dir / f"page_{page_num}.png"
                pix.save(str(img_path))

                try:
                    img  = Image.open(img_path)
                    img  = preprocess_image(img)
                    raw  = pytesseract.image_to_string(img, lang=lang, config="--oem 1 --psm 6")
                    text = clean_ocr_text(raw)
                except Exception as e:
                    text = f"[OCR Error: {e}]"

                (sec_dir / f"page_{page_num}.txt").write_text(text, encoding="utf-8")
                sec_pages.append({"page": page_num, "text": text})
                done += 1

                yield f"data: {json.dumps({'type':'page','section':name,'page':page_num,'done':done,'total':total_pages,'text':text})}\n\n"

            combined = "\n\n".join(f"{'='*60}\n  PAGE {p['page']}\n{'='*60}\n\n{p['text']}" for p in sec_pages)
            (sec_dir / "combined.txt").write_text(combined, encoding="utf-8")
            all_data.append({"name": name, "pages": sec_pages})
            yield f"data: {json.dumps({'type':'section_end','section':name})}\n\n"

        doc.close()
        (job_dir / "sections_data.json").write_text(json.dumps(all_data, ensure_ascii=False), encoding="utf-8")
        yield f"data: {json.dumps({'type':'done','job_id':job_id})}\n\n"

    return Response(generate(), mimetype="text/event-stream",
                    headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})


# ── GET /download/<job_id>/<fmt> ───────────────────────────────────────────────
@app.route("/download/<job_id>/<fmt>")
def download(job_id, fmt):
    sj = WORKSPACE / job_id / "sections_data.json"
    if not sj.exists():
        return jsonify({"error": "No output found"}), 404
    data = json.loads(sj.read_text(encoding="utf-8"))

    if fmt == "docx":
        if not DOCX_AVAILABLE:
            return jsonify({"error": "python-docx not installed"}), 500
        return send_file(io.BytesIO(build_docx(data)),
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True, download_name="extracted_text.docx")
    elif fmt == "zip":
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
            for sec in data:
                safe = _safe_dirname(sec['name'])
                zf.writestr(f"{safe}/combined.txt",
                    "\n\n".join(f"{'='*60}\n  PAGE {p['page']}\n{'='*60}\n\n{p['text']}" for p in sec['pages']))
                for p in sec['pages']:
                    zf.writestr(f"{safe}/page_{p['page']}.txt", p['text'])
        buf.seek(0)
        return send_file(buf, mimetype="application/zip", as_attachment=True, download_name="extracted_sections.zip")
    return jsonify({"error": "Unknown format"}), 400


if __name__ == "__main__":
    print("Starting PDF OCR Tool v3...")
    print("Open at: http://localhost:5050")
    app.run(host="0.0.0.0", port=5050, debug=False, threaded=True)
